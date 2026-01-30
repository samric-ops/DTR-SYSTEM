import streamlit as st
from io import BytesIO
from docx import Document
from calendar import monthrange
import datetime as dt

st.set_page_config(page_title="DTR Generator (Month-agnostic .docx)", page_icon="🗓️", layout="centered")
st.title("DTR Generator — Month-agnostic (D01..D31) with tokens")

st.markdown("Gamitin ang template na ito: **DTR_TOKENS_TABLE.docx** (2 halves, may D01..D31 tokens).")

# Sidebar: header fields
with st.sidebar:
    st.header("Header fields")
    emp_name = st.text_input("Name", "")
    emp_no = st.text_input("Employee No.", "")
    month_name = st.text_input("Month (text)", "")
    year_num = st.number_input("Year", min_value=1900, max_value=2100, value=2026, step=1)
    am_sched = st.text_input("AM schedule", "07:30 AM – 11:50 AM")
    pm_sched = st.text_input("PM schedule", "12:50 PM – 04:30 PM")
    sat_sched = st.text_input("Saturday schedule", "AS REQUIRED")

st.subheader("1) Template")
use_bundled = st.toggle("Gamitin ang bundled tokenized template", value=True)
if not use_bundled:
    uploaded = st.file_uploader("Upload .docx (DTR_TOKENS_TABLE)", type=["docx"])
else:
    uploaded = open("DTR_TOKENS_TABLE.docx", "rb")

st.subheader("2) Piliin ang buwan para sa UI (optional)")
colA, colB = st.columns(2)
with colA:
    year_ui = st.number_input("Year (UI)", min_value=1900, max_value=2100, value=int(year_num), step=1)
with colB:
    month_ui = st.number_input("Month # (UI)", min_value=1, max_value=12, value=1, step=1)

show_weekdays_only = st.checkbox("Ipakita lang ang weekdays sa form (Mon–Fri)", value=False)

# Build mapping base
mapping = {
    "{{NAME}}": emp_name,
    "{{EMP_NO}}": emp_no,
    "{{MONTH}}": month_name,
    "{{YEAR}}": str(year_num),
    "{{AM_SCHED}}": am_sched,
    "{{PM_SCHED}}": pm_sched,
    "{{SAT_SCHED}}": sat_sched,
}

# Initialize day tokens blank
for d in range(1,32):
    mapping[f"{{{{D{d:02d}_AM_IN}}}}"] = ""
    mapping[f"{{{{D{d:02d}_AM_OUT}}}}"] = ""
    mapping[f"{{{{D{d:02d}_PM_IN}}}}"] = ""
    mapping[f"{{{{D{d:02d}_PM_OUT}}}}"] = ""

if uploaded:
    st.subheader("3) Daily Time Entries")
    days = list(range(1, monthrange(year_ui, month_ui)[1]+1))
    if show_weekdays_only:
        days = [d for d in days if dt.date(year_ui, month_ui, d).weekday() < 5]

    st.caption("Maglagay ng oras (HH:MM / HH:MM AM/PM, depende sa gusto mo). Iwanang blanko ang mga hindi kailangang araw.")
    for d in days:
        c1,c2,c3,c4,c5 = st.columns([0.5,1,1,1,1])
        c1.markdown(f"**{d:02d}**")
        mapping[f"{{{{D{d:02d}_AM_IN}}}}"]  = c2.text_input(f"AM IN {d:02d}", key=f"ami{d}")
        mapping[f"{{{{D{d:02d}_AM_OUT}}}}"] = c3.text_input(f"AM OUT {d:02d}", key=f"amo{d}")
        mapping[f"{{{{D{d:02d}_PM_IN}}}}"]  = c4.text_input(f"PM IN {d:02d}", key=f"pmi{d}")
        mapping[f"{{{{D{d:02d}_PM_OUT}}}}"] = c5.text_input(f"PM OUT {d:02d}", key=f"pmo{d}")

    if st.button("Generate .docx"):
        doc = Document(uploaded)
        def replace_tokens(document, tokens):
            # Replace in body paragraphs
            for p in document.paragraphs:
                for r in p.runs:
                    if r.text:
                        for k, v in tokens.items():
                            if k in r.text:
                                r.text = r.text.replace(k, v or "")
            # Replace in tables
            for tbl in document.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for r in p.runs:
                                if r.text:
                                    for k, v in tokens.items():
                                        if k in r.text:
                                            r.text = r.text.replace(k, v or "")
        replace_tokens(doc, mapping)
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        st.success("Handa na ang DTR file!")
        st.download_button("⬇️ Download DTR_filled.docx", data=buf, file_name="DTR_filled.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
